#!/usr/bin/env python3
"""Convert main_google_docs.md into a .docx with embedded figures and real tables."""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(BASE, "main_google_docs.md")
OUT = os.path.join(BASE, "main_google_docs.docx")
FIG_DIRS = [os.path.join(BASE, "figures_png"), os.path.join(BASE, "figures")]

FIG_RE = re.compile(r"^\[FIGURE\s+\d+:\s*([^\s—]+)")


def resolve_png(name):
    base = os.path.splitext(os.path.basename(name))[0]
    for d in FIG_DIRS:
        p = os.path.join(d, base + ".png")
        if os.path.exists(p):
            return p
    return None


def add_runs(par, text):
    """Add inline runs handling **bold**, *italic*, `code`."""
    # tokenize on the three markers
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"
        else:
            r = par.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def main():
    with open(MD, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            i += 1
            continue

        # headings
        if stripped.startswith("#"):
            m = re.match(r"(#{1,6})\s+(.*)", stripped)
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            i += 1
            continue

        # figure placeholder
        fm = FIG_RE.match(stripped)
        if fm:
            png = resolve_png(fm.group(1))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if png:
                run = p.add_run()
                run.add_picture(png, width=Inches(6.0))
            else:
                add_runs(p, stripped)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.style = "Intense Quote"
            add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # table block
        if stripped.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            # drop separator rows (---)
            rows = []
            for r in block:
                cells = [c.strip() for c in r.strip("|").split("|")]
                if all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncol)
                tbl.style = "Light Grid Accent 1"
                for ri, r in enumerate(rows):
                    cells = tbl.add_row().cells
                    for ci in range(ncol):
                        txt = r[ci] if ci < len(r) else ""
                        cells[ci].text = ""
                        para = cells[ci].paragraphs[0]
                        add_runs(para, txt)
                        if ri == 0:
                            for run in para.runs:
                                run.bold = True
            continue

        # normal paragraph (single line; md here keeps paragraphs on one line)
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
